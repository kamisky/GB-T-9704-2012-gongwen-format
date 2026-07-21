# -*- coding: utf-8 -*-
"""生成后自检: 对照 GB/T 9704—2012 检查 .docx, 打印 PASS/FAIL, 偏差追加到 known_issues.md。
用法: python verify.py <公文.docx>
也可被 import: from verify import verify_docx
"""
import os, sys, datetime
from docx import Document
from docx.shared import Mm, Pt
from docx.oxml.ns import qn

KI = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                  "references", "known_issues.md")


def mm(v):
    """python-docx Length 是 EMU 的子类; 1mm=36000EMU。直接用 .mm 属性更稳。"""
    if v is None:
        return 0
    if hasattr(v, "mm"):
        return round(v.mm, 1)
    return round(v / 36000.0, 1)


def verify_docx(path):
    """检查 .docx 文件是否符合 GB/T 9704—2012 国标。
    
    Returns:
        dict: {
            "all_pass": bool,
            "total": int,
            "passed": int,
            "checks": [{"name": str, "ok": bool, "detail": str}, ...],
            "failures": [{"name": str, "detail": str}, ...],
        }
    """
    doc = Document(path)
    s = doc.sections[0]
    checks = []

    def check(name, ok, detail=""):
        checks.append({"name": name, "ok": bool(ok), "detail": detail})

    # ---- 页面 ----
    check("A4幅面 210×297",
          abs(mm(s.page_width) - 210) < 1 and abs(mm(s.page_height) - 297) < 1,
          f"{mm(s.page_width)}×{mm(s.page_height)}")
    check("天头37mm", abs(mm(s.top_margin) - 37) <= 1, f"{mm(s.top_margin)}")
    check("下边距35mm", abs(mm(s.bottom_margin) - 35) <= 1, f"{mm(s.bottom_margin)}")
    check("订口28mm", abs(mm(s.left_margin) - 28) <= 1, f"{mm(s.left_margin)}")
    check("切口26mm", abs(mm(s.right_margin) - 26) <= 1, f"{mm(s.right_margin)}")
    check("版心156×225",
          abs(mm(s.page_width - s.left_margin - s.right_margin) - 156) <= 2 and
          abs(mm(s.page_height - s.top_margin - s.bottom_margin) - 225) <= 2,
          f"{mm(s.page_width - s.left_margin - s.right_margin)}×{mm(s.page_height - s.top_margin - s.bottom_margin)}")

    # ---- Normal 样式 ----
    st = doc.styles["Normal"]
    sz = st.font.size.pt if st.font.size else 0
    check("默认3号(16pt)仿宋", abs(sz - 16) < 0.5, f"{sz}pt")

    # ---- 段落要素扫描 ----
    texts = [(p.text, p) for p in doc.paragraphs]
    joined = "\n".join(t for t, _ in texts)

    # 发文字号
    dn = next((t for t, _ in texts
               if (u"〔" in t and u"〕" in t and "印章" not in t and "加盖" not in t)
               or (t.strip().startswith("第") and "号" in t)), "")
    ok_dn = (u"〔" in dn and u"〕" in dn) or (dn.strip().startswith("第") and "号" in dn)
    check("发文字号合规(〔〕或第X号)", ok_dn, dn or "(未找到发文字号)")

    # 标题
    title_found = any(t for t, _ in texts if t.strip())
    check("存在标题段落", title_found, "")

    # 成文日期
    date_ok = any(("年" in t and "月" in t and "日" in t) for t, _ in texts)
    check("成文日期含年月日", date_ok, "")

    # 抄送格式
    copy_found = any("抄送：" in t for t, _ in texts)
    check("抄送格式(抄送：...。)",
          (not copy_found) or any(
              t.startswith("抄送：") and t.rstrip().endswith("。") for t, _ in texts),
          "" if not copy_found else "(检查末尾句号)")

    # ---- 页码 ----
    odd_xml = s.footer._element.xml
    even_xml = s.even_page_footer._element.xml
    check("页码含一字线(—)", "—" in odd_xml and "—" in even_xml, "")
    check("页码含PAGE域", "PAGE" in odd_xml.upper() and "PAGE" in even_xml.upper(), "")
    check("开启奇偶页不同",
          s.footer.is_linked_to_previous is False and
          s.even_page_footer.is_linked_to_previous is False, "")

    # ---- 红色分隔线 ----
    has_red_border = any("FF0000" in p._element.xml.upper() for p in doc.paragraphs)
    check("版头红色分隔线", has_red_border, "")

    # ---- 版记分隔线 ----
    has_banji_text = ("抄送" in joined) or ("印发" in joined)
    has_banji_border = any('w:color="000000"' in p._element.xml for p in doc.paragraphs)
    if has_banji_text:
        check("版记分隔线(黑色)", has_banji_border, "")

    # ---- 汇总 ----
    failures = [c for c in checks if not c["ok"]]
    return {
        "all_pass": len(failures) == 0,
        "total": len(checks),
        "passed": len(checks) - len(failures),
        "checks": checks,
        "failures": failures,
    }


def print_report(result, path=""):
    """打印自检报告。"""
    print("=" * 50)
    print("GB/T 9704—2012 自检报告")
    print("=" * 50)
    for c in result["checks"]:
        tag = "PASS" if c["ok"] else "FAIL"
        detail = f"  ({c['detail']})" if c["detail"] else ""
        print(f"[{tag}] {c['name']}{detail}")
    print("-" * 50)
    print(f"通过 {result['passed']}/{result['total']}")

    if result["failures"]:
        with open(KI, "a", encoding="utf-8") as f:
            f.write(f"\n## [{datetime.date.today()}] 自检偏差 ({os.path.basename(path) if path else '未知'})\n")
            for c in result["failures"]:
                f.write(f"- 现象: {c['name']} 不合规" +
                        (f" / {c['detail']}" if c["detail"] else "") + "\n")
                f.write(f"- 国标依据: 见 references/gbt9704-2012-spec.md\n")
                f.write(f"- 修复: 待核实(可能为渲染/近似或脚本逻辑)\n")
        print(f"偏差已追加至: {KI}")


def main():
    if len(sys.argv) < 2:
        print("用法: python verify.py <公文.docx>")
        return 2
    path = os.path.abspath(sys.argv[1])
    result = verify_docx(path)
    print_report(result, path)
    return 1 if result["failures"] else 0


if __name__ == "__main__":
    sys.exit(main())