# -*- coding: utf-8 -*-
"""GB/T 9704—2012 公文格式 Skill 生成器。
用法: python gongwen_gen.py <content.json> [输出.docx]
content.json 结构见 templates/content_schema.json。覆盖普通文件/信函/命令(令)/纪要四种格式。
参数全部来自 references/gbt9704-2012-spec.md。
"""
import os, sys, json

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ---- 字体检测 ----
_HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, _HERE)
try:
    from font_check import resolve_fonts
except Exception:
    resolve_fonts = None

# ============================================================
# 国标常量
# ============================================================
PAGE_W, PAGE_H = 210, 297          # A4 mm
M_TOP, M_BOT, M_LEFT, M_RIGHT = 37, 35, 28, 26   # 天头/下/订口/切口 mm
BANXIN_W, BANXIN_H = 156, 225      # 版心 mm
LINE_PT = 28.99                    # 22行撑满225mm: 225/22*2.8346
SIZE_BODY   = 16   # 三号
SIZE_TITLE  = 22   # 二号
SIZE_HEADER = 36   # 小初(发文机关标志)
SIZE_FOUR   = 14   # 四号(抄送/印发/页码)
RED = "FF0000"
BLACK = "000000"

# 层次字体映射
LAYER_FONT = {"h1": "heiti", "h2": "kaiti", "h3": "fangsong", "h4": "fangsong"}


def _fonts():
    if resolve_fonts:
        return resolve_fonts()
    return {"xiaobiaosong": "宋体", "fangsong": "仿宋", "kaiti": "楷体",
            "heiti": "黑体", "songti": "宋体"}


# ============================================================
# OOXML 辅助
# ============================================================
def set_run_font(run, eastasia, size_pt, bold=False, color=None):
    run.font.size = Pt(size_pt)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    for tag in ("rFonts", "sz", "szCs"):
        for ex in rpr.findall(qn(f"w:{tag}")):
            rpr.remove(ex)
    rpr.append(parse_xml(
        f'<w:rFonts {nsdecls("w")} w:ascii="{eastasia}" w:hAnsi="{eastasia}" '
        f'w:eastAsia="{eastasia}" w:cs="{eastasia}"/>'))
    rpr.append(parse_xml(f'<w:sz {nsdecls("w")} w:val="{int(size_pt*2)}"/>'))
    rpr.append(parse_xml(f'<w:szCs {nsdecls("w")} w:val="{int(size_pt*2)}"/>'))


def set_spacing(para, line_pt=LINE_PT, before_pt=0, after_pt=0, rule="exact"):
    ppr = para._element.get_or_add_pPr()
    for ex in ppr.findall(qn("w:spacing")):
        ppr.remove(ex)
    a = f'w:line="{int(round(line_pt*20))}" w:lineRule="{rule}"'
    if before_pt:
        a += f' w:before="{int(round(before_pt*20))}" w:beforeLines="0"'
    if after_pt:
        a += f' w:after="{int(round(after_pt*20))}" w:afterLines="0"'
    ppr.append(parse_xml(f'<w:spacing {nsdecls("w")} {a}/>'))


def set_indents(para, first_chars=0, left_chars=0, right_chars=0):
    ppr = para._element.get_or_add_pPr()
    for ex in ppr.findall(qn("w:ind")):
        ppr.remove(ex)
    a = ""
    if first_chars:
        a += f' w:firstLineChars="{int(first_chars*100)}" w:firstLine="{int(first_chars*SIZE_BODY*20)}"'
    if left_chars:
        a += f' w:leftChars="{int(left_chars*100)}" w:left="{int(left_chars*SIZE_BODY*20)}"'
    if right_chars:
        a += f' w:rightChars="{int(right_chars*100)}" w:right="{int(right_chars*SIZE_BODY*20)}"'
    if a:
        ppr.append(parse_xml(f'<w:ind {nsdecls("w")} {a.strip()}/>'))


def add_border(para, edge="bottom", color=RED, sz=12, space=4, val="single"):
    ppr = para._element.get_or_add_pPr()
    pbdr = ppr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = parse_xml(f'<w:pBdr {nsdecls("w")}></w:pBdr>')
        ppr.append(pbdr)
    for ex in pbdr.findall(qn(f"w:{edge}")):
        pbdr.remove(ex)
    pbdr.append(parse_xml(
        f'<w:{edge} {nsdecls("w")} w:val="{val}" w:sz="{sz}" '
        f'w:space="{space}" w:color="{color}"/>'))


def add_page_field(paragraph):
    run = paragraph.add_run()
    run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    run._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    return run


def enable_even_odd(doc):
    """开启奇偶页不同(用于页码奇右偶左)。"""
    settings = doc.settings.element
    for ex in settings.findall(qn("w:evenAndOddHeaders")):
        settings.remove(ex)
    settings.append(parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")} w:val="true"/>'))


# ============================================================
# 构建器
# ============================================================
class GongwenBuilder:
    def __init__(self, content):
        self.c = content
        self.f = _fonts()
        self.doc = Document()
        self._setup_page()
        self._setup_normal()

    # ---- 页面 ----
    def _setup_page(self):
        s = self.doc.sections[0]
        s.page_width, s.page_height = Mm(PAGE_W), Mm(PAGE_H)
        s.top_margin, s.bottom_margin = Mm(M_TOP), Mm(M_BOT)
        s.left_margin, s.right_margin = Mm(M_LEFT), Mm(M_RIGHT)
        s.footer_distance = Mm(25)   # 页码约在版心下边缘下7mm(近似)

    def _setup_normal(self):
        st = self.doc.styles["Normal"]
        st.font.name = self.f["fangsong"]
        st.font.size = Pt(SIZE_BODY)
        st.element.rPr.rFonts.set(qn("w:eastAsia"), self.f["fangsong"])
        pf = st.paragraph_format
        pf.line_spacing = Pt(LINE_PT)

    # ---- 基本段落工厂 ----
    def _para(self, align=None, line=LINE_PT, before=0, after=0,
              first_chars=0, left_chars=0, right_chars=0):
        p = self.doc.add_paragraph()
        if align is not None:
            p.alignment = align
        set_spacing(p, line, before, after)
        set_indents(p, first_chars, left_chars, right_chars)
        return p

    def _run(self, para, text, font_key, size, bold=False, color=None):
        r = para.add_run(text)
        set_run_font(r, self.f[font_key], size, bold, color)
        return r

    # ================= 版头 =================
    def build_biaotou(self):
        c = self.c
        fmt = c.get("format", "standard")
        # 份号/密级/紧急 (普通格式: 左上角)
        if fmt == "standard":
            n_top = 0
            if c.get("fenhao"):
                p = self._para(left_chars=0); self._run(p, str(c["fenhao"]).zfill(6), "fangsong", SIZE_BODY); n_top += 1
            if c.get("secret_level"):
                p = self._para(left_chars=0)
                t = c["secret_level"] + ("★" + c.get("secret_period", "") if c.get("secret_period") else "")
                # 保密期限写法: 密级★保密期限
                t = c["secret_level"] + ("★" + c["secret_period"] if c.get("secret_period") else "")
                self._run(p, t, "heiti", SIZE_BODY); n_top += 1
            if c.get("urgency"):
                p = self._para(left_chars=0); self._run(p, c["urgency"], "heiti", SIZE_BODY); n_top += 1
        else:
            n_top = 0
        # 发文机关标志 / 纪要标志
        self._issuer_mark(fmt, n_top)
        # 发文字号 + 签发人 + 红线
        self._doc_number_row(fmt)
        # 红色分隔线(发文字号下4mm, 与版心等宽)
        self._separator(RED, sz=12, gap_mm=4)

    def _issuer_mark(self, fmt, n_top):
        c = self.c
        xbs = self.f["xiaobiaosong"]
        is_fb = (xbs == self.f["songti"])   # 是否回退宋体
        if fmt == "standard":
            text = c.get("issuer", "") + ("文件" if c.get("issuer") and "文件" not in c.get("issuer", "") and c.get("add_wenjian", True) else "")
            # 若 issuer 已含"文件"则不重复
            if c.get("issuer") and c["issuer"].endswith("文件"):
                text = c["issuer"]
            size = SIZE_HEADER
            top_mm = 35
        elif fmt == "letter":
            text = c.get("issuer", "")
            size = SIZE_HEADER
            top_mm = 30  # 上边缘至上页边30mm
        elif fmt == "order":
            text = c.get("issuer", "") + ("命令" if not c.get("issuer", "").endswith(("命令", "令")) else "")
            size = SIZE_HEADER
            top_mm = 20  # 至版心上边缘20mm
        elif fmt == "minutes":
            text = c.get("issuer", "") + "纪要"
            size = SIZE_HEADER
            top_mm = 35
        else:
            text = c.get("issuer", ""); size = SIZE_HEADER; top_mm = 35
        # 顶部留白: top_mm 减去已占行高
        consumed_mm = n_top * LINE_PT * 0.3528
        before_mm = max(0, top_mm - consumed_mm)
        p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, before=before_mm*2.8346,
                       line=size*1.2)
        self._run(p, text, "xiaobiaosong", size, bold=is_fb, color=RED)

    def _doc_number_row(self, fmt):
        c = self.c
        if fmt == "order":
            # 令号: 标志下空二行居中
            if c.get("doc_number"):
                p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, before=2*LINE_PT)
                self._run(p, c["doc_number"], "fangsong", SIZE_BODY)
            return
        if not c.get("doc_number") and not c.get("signer"):
            # 无发文字号也需空二行后再红线
            self._para(before=2*LINE_PT)
            return
        p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, before=2*LINE_PT)
        if c.get("signer"):
            # 上行文: 发文字号居左空一字 + 签发人居右空一字(同一行)
            set_indents(p, left_chars=1, right_chars=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if c.get("doc_number"):
                self._run(p, c["doc_number"], "fangsong", SIZE_BODY)
            # 右对齐签发人
            ts = p.paragraph_format.tab_stops
            ts.add_tab_stop(Mm(BANXIN_W - 2*SIZE_BODY*0.3528), WD_TAB_ALIGNMENT.RIGHT)
            p.add_run("\t")
            self._run(p, "签发人：", "fangsong", SIZE_BODY)
            self._run(p, c["signer"], "kaiti", SIZE_BODY)
        else:
            if c.get("doc_number"):
                self._run(p, c["doc_number"], "fangsong", SIZE_BODY)

    def _separator(self, color, sz=12, gap_mm=4):
        # 与版心等宽的整行分隔线: 空段落+底线, 行高=gap_mm
        p = self._para(line=gap_mm*2.8346)
        set_indents(p, 0, 0, 0)
        r = p.add_run("")
        set_run_font(r, self.f["fangsong"], SIZE_BODY)
        add_border(p, "bottom", color=color, sz=sz, space=1)

    def _double_separator(self, color, sz=12, gap_mm=4):
        # 红色双线(信函格式): w:val="double"
        p = self._para(line=gap_mm*2.8346)
        set_indents(p, 0, 0, 0)
        r = p.add_run("")
        set_run_font(r, self.f["fangsong"], SIZE_BODY)
        add_border(p, "bottom", color=color, sz=sz, space=1, val="double")

    # ================= 主体 =================
    def build_body(self):
        c = self.c
        # 标题
        if c.get("title"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.CENTER, before=2*LINE_PT, line=SIZE_TITLE*1.3)
            lines = c["title"].split("\n")
            for i, ln in enumerate(lines):
                if i:
                    p.add_run().add_break()
                self._run(p, ln, "xiaobiaosong", SIZE_TITLE,
                          bold=(self.f["xiaobiaosong"] == self.f["songti"]))
        # 主送机关
        if c.get("recipient"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, before=LINE_PT)
            self._run(p, c["recipient"], "fangsong", SIZE_BODY)
        # 正文
        for item in (c.get("body") or []):
            self._body_item(item)
        # 附件说明
        if c.get("attachments"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, before=LINE_PT, left_chars=2)
            atts = c["attachments"]
            if len(atts) == 1:
                self._run(p, "附件：" + atts[0]["name"], "fangsong", SIZE_BODY)
            else:
                self._run(p, "附件：", "fangsong", SIZE_BODY)
                for i, a in enumerate(atts, 1):
                    self._run(p, f"\n{i}. {a['name']}", "fangsong", SIZE_BODY)
        # 署名/日期/印章
        self._signature_block()
        # 附注
        if c.get("annotation"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, left_chars=2)
            ann = c["annotation"]
            if not (ann.startswith("(") or ann.startswith("（")):
                ann = "（" + ann + "）"
            self._run(p, ann, "fangsong", SIZE_BODY)
        # 纪要: 出席/请假/列席
        if c.get("format") == "minutes":
            self._minutes_lists()

    def _body_item(self, item):
        t = item.get("type", "para")
        txt = item.get("text", "")
        if t == "para":
            p = self._para(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_chars=2)
            self._run(p, txt, "fangsong", SIZE_BODY)
        elif t == "h1":
            p = self._para(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            self._run(p, txt, "heiti", SIZE_BODY)
        elif t == "h2":
            p = self._para(align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            self._run(p, txt, "kaiti", SIZE_BODY)
        elif t in ("h3", "h4"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_chars=2)
            self._run(p, txt, "fangsong", SIZE_BODY)

    def _signature_block(self):
        c = self.c
        seal = c.get("seal", "placeholder")
        if not (c.get("issuer_name") or c.get("date")):
            return
        # 发文机关署名(成文日期之上,以日期为准居中)
        if c.get("issuer_name"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.RIGHT, right_chars=4, before=LINE_PT)
            self._run(p, c["issuer_name"], "fangsong", SIZE_BODY)
        # 印章占位
        if seal == "placeholder" and (c.get("issuer_name") or c.get("date")):
            p = self._para(align=WD_ALIGN_PARAGRAPH.RIGHT, right_chars=4, line=SIZE_BODY*1.2)
            self._run(p, "〔此处加盖红色印章〕", "fangsong", SIZE_BODY, color="C00000")
        # 成文日期(右空四字)
        if c.get("date"):
            p = self._para(align=WD_ALIGN_PARAGRAPH.RIGHT, right_chars=4)
            self._run(p, c["date"], "fangsong", SIZE_BODY)

    def _minutes_lists(self):
        c = self.c
        for key, label in (("attendees", "出席"), ("leave", "请假"), ("observers", "列席")):
            if c.get(key):
                p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, before=LINE_PT, left_chars=2)
                self._run(p, label + "：", "heiti", SIZE_BODY)
                self._run(p, c[key], "fangsong", SIZE_BODY)

    # ================= 版记 =================
    def build_banji(self):
        c = self.c
        fmt = c.get("format", "standard")
        has_copy = bool(c.get("copy_to"))
        has_print = bool(c.get("printer") or c.get("print_date"))
        is_letter = (fmt == "letter")
        if is_letter:
            # 信函: 版记不加印发机关/日期/分隔线, 位于最下方
            if has_copy:
                p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, left_chars=1, right_chars=1)
                names = "、".join(c["copy_to"]) if isinstance(c["copy_to"], list) else c["copy_to"]
                self._run(p, "抄送：" + names + "。", "fangsong", SIZE_FOUR)
            return
        if not (has_copy or has_print):
            return
        # 首条粗线
        self._separator(BLACK, sz=8, gap_mm=2)
        if has_copy:
            p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, left_chars=1, right_chars=1)
            names = "、".join(c["copy_to"]) if isinstance(c["copy_to"], list) else c["copy_to"]
            text = "抄送：" + names + "。"
            self._run(p, text, "fangsong", SIZE_FOUR)
            if has_print:
                self._separator(BLACK, sz=6, gap_mm=2)   # 中间细线
        if has_print:
            p = self._para(align=WD_ALIGN_PARAGRAPH.LEFT, left_chars=1, right_chars=1)
            self._run(p, c.get("printer", "") + "    ", "fangsong", SIZE_FOUR)
            ts = p.paragraph_format.tab_stops
            ts.add_tab_stop(Mm(BANXIN_W - 2*SIZE_FOUR*0.3528), WD_TAB_ALIGNMENT.RIGHT)
            p.add_run("\t")
            d = c.get("print_date", "") + "印发"
            self._run(p, d, "fangsong", SIZE_FOUR)
        # 末条粗线(与版心下边缘重合)
        self._separator(BLACK, sz=8, gap_mm=2)

    # ================= 页码 =================
    def build_pagenum(self):
        c = self.c
        fmt = c.get("format", "standard")
        s = self.doc.sections[0]
        enable_even_odd(self.doc)
        # 首页页码: 信函格式首页不显示
        if fmt == "letter":
            s.different_first_page_header_footer = True
            ffp = s.first_page_footer
            ffp.is_linked_to_previous = False
            ffp.paragraphs[0].text = ""
        # 奇页(单页码)居右空一字
        odd = s.footer
        odd.is_linked_to_previous = False
        op = odd.paragraphs[0]
        op.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        set_indents(op, right_chars=1)
        set_spacing(op, line_pt=SIZE_FOUR)
        self._run(op, "— ", "songti", SIZE_FOUR)
        add_page_field(op)
        self._run(op, " —", "songti", SIZE_FOUR)
        # 偶页(双页码)居左空一字
        even = s.even_page_footer
        even.is_linked_to_previous = False
        ep = even.paragraphs[0]
        ep.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_indents(ep, left_chars=1)
        set_spacing(ep, line_pt=SIZE_FOUR)
        self._run(ep, "— ", "songti", SIZE_FOUR)
        add_page_field(ep)
        self._run(ep, " —", "songti", SIZE_FOUR)

    # ================= 主流程 =================
    def build(self):
        fmt = self.c.get("format", "standard")
        self.build_biaotou()
        if fmt == "letter":
            # 上粗下细红双线(标志下4mm)
            self._double_separator(RED, sz=12, gap_mm=4)
        self.build_body()
        if fmt == "letter":
            # 距下页边20mm红双线(上细下粗) - 近似: 置于版记之前
            self._double_separator(RED, sz=12, gap_mm=4)
        self.build_banji()
        self.build_pagenum()
        return self.doc


def main():
    if len(sys.argv) < 2:
        print("用法: python gongwen_gen.py <content.json> [输出.docx]")
        return 2
    src = os.path.abspath(sys.argv[1])
    with open(src, "r", encoding="utf-8") as f:
        content = json.load(f)
    out = os.path.abspath(sys.argv[2]) if len(sys.argv) > 2 else \
        os.path.splitext(src)[0] + ".docx"
    doc = GongwenBuilder(content).build()
    doc.save(out)
    print(f"[OK] 公文已生成: {out}")
    print(f"     格式: {content.get('format','standard')} | 字体: 小标宋={'已装' if _fonts().get('xiaobiaosong')!=_fonts().get('songti') else '回退宋体'}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
